# -*- coding: utf-8 -*-
"""任务模板章节配置 → 报告章节取舍（enabled_sections）"""
from docx import Document

from models import Inspection
from utils.report_generator import _section_enabled, generate_inspection_report_v4


class TestSectionEnabledHelper:
    def test_default_all_enabled_when_no_config(self):
        assert _section_enabled({}, 'topology') is True
        assert _section_enabled({'enabled_sections': None}, 'topology') is True

    def test_disabled_section_detected(self):
        sections = {'enabled_sections': [
            {'key': 'topology', 'title': '六、网络拓扑图', 'enabled': False},
            {'key': 'suggestions', 'title': '九、运行建议', 'enabled': True},
        ]}
        assert _section_enabled(sections, 'topology') is False
        assert _section_enabled(sections, 'suggestions') is True
        # 未列出的 key 默认启用
        assert _section_enabled(sections, 'photos') is True


class TestReportRespectsEnabledSections:
    def test_disabled_chapters_absent_from_docx(self, app, tmp_path, monkeypatch):
        """禁用章节的标题不出现在生成的 docx 中"""
        with app.app_context():
            insp = Inspection(title='章节测试巡检', customer_id=None)
            sections = {
                'q2_1': '季度巡检内容',
                'flood_advice': '运行建议内容',
                'tech_support': '13800000000',
                'enabled_sections': [
                    {'key': 'quarterly_work', 'enabled': True},
                    {'key': 'topology', 'enabled': False},      # 禁用拓扑章
                    {'key': 'suggestions', 'enabled': False},   # 禁用建议章
                    {'key': 'support_contacts', 'enabled': True},
                ],
            }
            monkeypatch.setattr('utils.report_generator.os.path.dirname',
                                lambda p: str(tmp_path) if 'report_generator' in p else p)
            path = generate_inspection_report_v4(insp, '测试客户', device_results=[], sections=sections)
            doc = Document(path)
            full_text = '\n'.join(p.text for p in doc.paragraphs)
            assert '季度运维工作内容' in full_text        # 启用 → 保留
            assert '季度巡检内容' in full_text
            assert '网络拓扑图' not in full_text          # 禁用 → 整章移除
            assert '运行建议内容' not in full_text        # 禁用 → 内容不出现
            assert '售后服务电话' in full_text            # 启用 → 保留
            assert '13800000000' in full_text

    def test_legacy_report_without_config_unchanged(self, app, tmp_path):
        """旧记录（无 enabled_sections）报告章节完整"""
        with app.app_context():
            insp = Inspection(title='旧格式巡检', customer_id=None)
            path = generate_inspection_report_v4(insp, '测试客户', device_results=[],
                                                 sections={'flood_advice': '建议X'})
            doc = Document(path)
            full_text = '\n'.join(p.text for p in doc.paragraphs)
            assert '网络拓扑图' in full_text
            assert '建议X' in full_text

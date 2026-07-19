"""Word 报告自动生成模块"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from flask import current_app


def _chinese_uppercase_date(d):
    """将日期转为中文大写，如 2026-05-21 → 二〇二六年五月二十一日"""
    digits = ['〇', '一', '二', '三', '四', '五', '六', '七', '八', '九']
    months = ['一月', '二月', '三月', '四月', '五月', '六月',
              '七月', '八月', '九月', '十月', '十一月', '十二月']
    # 日期的特殊写法
    def _day_str(day):
        if day < 1 or day > 31:
            return str(day)
        if day <= 10:
            if day == 10:
                return '十日'
            return digits[day] + '日'
        elif day < 20:
            prefix = '十'
            unit = day % 10
            if unit == 0:
                return '十日'
            return prefix + digits[unit] + '日'
        elif day == 20:
            return '二十日'
        elif day < 30:
            return '二十' + digits[day % 10] + '日'
        elif day == 30:
            return '三十日'
        else:
            return '三十一' if day == 31 else '三十' + digits[day % 10] + '日'

    year_str = ''.join(digits[int(ch)] for ch in f'{d.year:04d}') + '年'
    month_str = months[d.month - 1]
    day_str = _day_str(d.day)
    return year_str + month_str + day_str


def _set_run_font(run, font_name='宋体', size_pt=10, bold=False, color=None):
    """设置 run 的字体属性"""
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


def _set_cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    """设置表格单元格文本"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold:
        run.bold = True


def _add_table_row(table, cells_data, bold=False, header=False):
    """添加表格行"""
    row = table.add_row()
    for i, val in enumerate(cells_data):
        if i < len(row.cells):
            _set_cell_text(row.cells[i], val, bold=bold)
    return row


def _add_chapter_heading(doc, number, title):
    """添加章节标题（Heading 2 样式，宋体三号16pt，黑色）"""
    p = doc.add_heading(f'{number}、{title}', level=2)
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(16)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def _add_sub_heading(doc, num, title):
    """添加子章节标题（Heading 3 样式，宋体四号14pt加粗，黑色）"""
    p = doc.add_heading(f'{num} {title}', level=3)
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def _add_toc(doc, chapters):
    """添加目录页（Word 自动目录 TOC 域）"""
    from docx.oxml import OxmlElement
    for _ in range(3):
        doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run('目  录')
    _set_run_font(run, font_name='宋体', size_pt=22, bold=True)
    doc.add_paragraph()

    # 插入 Word TOC 域代码（基于 Heading 1-3 样式）
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))

    def _add_toc_field(prg, code, text):
        r1 = prg.add_run()
        fc1 = OxmlElement('w:fldChar')
        fc1.set(qn('w:fldCharType'), 'begin')
        r1._element.append(fc1)
        r2 = prg.add_run()
        it = OxmlElement('w:instrText')
        it.set(qn('xml:space'), 'preserve')
        it.text = code
        r2._element.append(it)
        r3 = prg.add_run()
        fc2 = OxmlElement('w:fldChar')
        fc2.set(qn('w:fldCharType'), 'separate')
        r3._element.append(fc2)
        r4 = prg.add_run(text)
        r4.font.color.rgb = RGBColor(128, 128, 128)
        r4.font.size = Pt(12)
        r5 = prg.add_run()
        fc3 = OxmlElement('w:fldChar')
        fc3.set(qn('w:fldCharType'), 'end')
        r5._element.append(fc3)

    _add_toc_field(p, r' TOC \o "1-3" \h \z \u ', '【右键此处 → 更新域 → 更新整个目录，页码自动生成】')
    doc.add_page_break()


def _add_section_photos(doc, sections, key):
    """在章节中插入照片"""
    photos = sections.get(f'{key}_photos', [])
    if photos:
        for photo_path in photos:
            full_path = os.path.normpath(os.path.join(os.path.dirname(os.path.dirname(__file__)), photo_path))
            if os.path.exists(full_path):
                try:
                    doc.add_picture(full_path, width=Inches(3.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass


# ============================
# V4: 9章固定结构巡检报告生成器
# ============================

def generate_inspection_report_v4(inspection, customer_name, device_results=None, sections=None):
    """按固定9章结构生成巡检报告（V4复合巡检）。参见模块文档了解完整参数说明。"""
    import json

    if sections is None:
        sections = {}
    if device_results is None:
        device_results = []
    if not device_results and inspection:
        try:
            fv = json.loads(inspection.field_values_json) if isinstance(inspection.field_values_json, str) else (inspection.field_values_json or {})
            device_results = _build_device_results_from_values(inspection, fv)
        except:
            pass

    doc = Document()
    date_cn = _chinese_uppercase_date(
        inspection.inspection_date if hasattr(inspection, 'inspection_date') and inspection.inspection_date and hasattr(inspection.inspection_date, 'year') else datetime.now().date()
    )

    # 封面
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{customer_name}网络巡检运维服务报告')
    _set_run_font(run, '宋体', 28, bold=True)
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('江西丰功信息技术有限公司')
    _set_run_font(run2, '宋体', 22, bold=True)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(date_cn)
    _set_run_font(run3, '宋体', 18)
    doc.add_page_break()

    # 目录
    _add_toc(doc, ['总体情况', '季度运维工作内容', '巡检记录表', '现场图片', '故障工单', '网络拓扑图', '设备台账', '网络运行建议', '售后服务电话'])

    # === Ch1: 总体情况 ===
    _add_chapter_heading(doc, '一', '总体情况')
    doc.add_paragraph(f'本次机房巡检工作于{date_cn}全面展开。')
    _add_sub_heading(doc, '1.1', '运维对象')
    doc.add_paragraph('▸ 机房')
    if device_results:
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        _add_table_row(tbl, ['序号', '项目', 'IP地址', '管辖', '运行态势'], bold=True, header=True)
        for idx, dr in enumerate(device_results, 1):
            status = '运行良好'
            for item in dr.get('items', []):
                if item.get('value', '') in ('不正常', '异常', '低于标准', '已过期'):
                    status = '故障'
                    break
            _add_table_row(tbl, [str(idx), dr.get('device_name', '-'), dr.get('ip_address', '-'), customer_name, status])
        doc.add_paragraph()

    _add_sub_heading(doc, '1.2', '防汛机房情况汇报')
    doc.add_paragraph(sections.get('q1_2', ''))

    _add_sub_heading(doc, '1.3', '现场巡检隐患清单')
    doc.add_paragraph('根据现场巡检及日常运维发现以下隐患：')
    abnormal_items = []
    for dr in device_results:
        for item in dr.get('items', []):
            val = item.get('value', '')
            if val in ('不正常', '异常', '低于标准', '已过期', '即将到期'):
                abnormal_items.append({'cat': dr.get('template_category', ''), 'dev': dr.get('device_name', ''), 'item': item.get('name', ''), 'status': val})
    if abnormal_items:
        tbl2 = doc.add_table(rows=1, cols=5)
        tbl2.style = 'Table Grid'
        tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
        _add_table_row(tbl2, ['序号', '项目', '子项', '巡检情况', '备注'], bold=True, header=True)
        for idx, a in enumerate(abnormal_items, 1):
            _add_table_row(tbl2, [str(idx), a['cat'], f'{a["dev"]}-{a["item"]}', a['status'], ''])
    else:
        doc.add_paragraph('本次巡检未发现安全隐患。')

    _add_sub_heading(doc, '1.4', '隐患详情及影响')
    try:
        sr = json.loads(inspection.skip_reasons_json) if isinstance(inspection.skip_reasons_json, str) else (inspection.skip_reasons_json or {})
    except Exception as _e:
        current_app.logger.warning('解析 skip_reasons_json 失败：%s', repr(_e))
        sr = {}
    issue_idx = 1
    for dev_key, dev_sr in sr.items():
        if isinstance(dev_sr, dict):
            for item_name, reason_info in dev_sr.items():
                reason = reason_info.get('reason', '') if isinstance(reason_info, dict) else str(reason_info)
                detail = reason_info.get('detail', '') if isinstance(reason_info, dict) else ''
                if reason:
                    doc.add_paragraph(f'{issue_idx}、{dev_key} -- {item_name}')
                    doc.add_paragraph(f'   隐患内容：{reason}')
                    if detail:
                        doc.add_paragraph(f'   详情：{detail}')
                    issue_idx += 1
    if issue_idx == 1:
        doc.add_paragraph('本次巡检未发现重大隐患。')

    # === Ch2: 季度运维工作 ===
    _add_chapter_heading(doc, '二', '季度运维工作内容')
    for sub, title in [('q2_1', '季度巡检'), ('q2_2', '机房环境检查'), ('q2_3', '网络设备配置备份'),
                        ('q2_4', '标签及线缆检查'), ('q2_5', '核心交换机空接口做shutdown处理')]:
        _add_sub_heading(doc, sub.replace('q', '').replace('_', '.'), title)
        doc.add_paragraph(sections.get(sub, ''))
        if sub in ('q2_3', 'q2_4'):
            _add_section_photos(doc, sections, sub.replace('q', ''))

    # === Ch3: 巡检记录表 ===
    _add_chapter_heading(doc, '三', '巡检记录表')
    cat_order = {'服务器': 0, '网络设备': 1, '安全设备': 2, 'UPS': 3, '空调': 4}
    sorted_devices = sorted(device_results, key=lambda d: cat_order.get(d.get('template_category', ''), 99))
    counters = {}
    base_nums = {'服务器': '3.1', '网络设备': '3.2', '安全设备': '3.5', 'UPS': '3.12', '空调': '3.13'}

    for dr in sorted_devices:
        cat = dr.get('template_category', '其他')
        counters[cat] = counters.get(cat, 0) + 1
        idx = counters[cat]
        base = base_nums.get(cat, '3.1')
        try:
            parts = base.split('.')
            sec_num = f'{parts[0]}.{int(parts[1]) + idx - 1}'
        except:
            sec_num = f'3.{idx}'
        _add_sub_heading(doc, f'{sec_num}．', f'{cat}巡检记录表-{idx}')

        # 设备信息头表
        info_t = doc.add_table(rows=4, cols=2)
        info_t.style = 'Table Grid'
        info_t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for (k1, v1, k2, v2), row_idx in zip([
            ('设备名称', dr.get('device_name', '-'), '安装位置', dr.get('location', '-')),
            ('设备型号', dr.get('model', '-'), '管理IP地址', dr.get('ip_address', '-')),
            ('操作系统版本', dr.get('os_version', '-'), '系统运行时间', dr.get('uptime', '-')),
            ('匹配模板', dr.get('template_name', '-'), '设备类型', dr.get('device_type', '-')),
        ], range(4)):
            _set_cell_text(info_t.rows[row_idx].cells[0], f'{k1}: {v1}')
            _set_cell_text(info_t.rows[row_idx].cells[1], f'{k2}: {v2}')
        doc.add_paragraph()

        # 检查项表（5列）
        items = dr.get('items', [])
        if items:
            chk_t = doc.add_table(rows=1, cols=5)
            chk_t.style = 'Table Grid'
            chk_t.alignment = WD_TABLE_ALIGNMENT.CENTER
            _add_table_row(chk_t, ['序号', '巡检内容', '检查项目说明', '巡检情况说明', '备注'], bold=True, header=True)
            for item_idx, item in enumerate(items, 1):
                # V14: 子项目以「主项目.子项目标签」形式存储 → 美化为「主项目 · 子项目」
                raw_name = item.get('name', '')
                name = raw_name.replace('.', ' · ') if '.' in raw_name else raw_name
                help_txt = item.get('help_text', '')
                val = item.get('value', '-')
                ft = item.get('field_type', 'text')
                if ft == 'dropdown':
                    if val in ('正常', '符合要求'):
                        disp = '☑ 正常 □ 其他'
                    else:
                        disp = f'□ 正常 ☑ {val}'
                elif ft == 'config_backup':
                    disp = f'已上传：{val}' if val and val != '-' else '未上传'
                else:
                    disp = str(val)
                _add_table_row(chk_t, [str(item_idx), name, help_txt, disp, ''])
        doc.add_paragraph()

    # === Ch4-9 ===
    _add_chapter_heading(doc, '四', '现场图片')
    doc.add_paragraph(sections.get('q4_images', '（暂无现场图片）'))
    _add_chapter_heading(doc, '五', '故障工单')
    doc.add_paragraph('无')
    _add_chapter_heading(doc, '六', '网络拓扑图')
    topology_photos = sections.get('topology_photos', [])
    if topology_photos:
        for tp in topology_photos:
            if os.path.exists(tp.lstrip('/')):
                try:
                    doc.add_picture(tp.lstrip('/'), width=Inches(5.0))
                except:
                    pass
    _add_chapter_heading(doc, '七', '设备台账')
    if device_results:
        ldr_t = doc.add_table(rows=1, cols=5)
        ldr_t.style = 'Table Grid'
        ldr_t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _add_table_row(ldr_t, ['序号', '设备名称', '设备型号', '管理IP', '位置'], bold=True, header=True)
        for idx, dr in enumerate(device_results, 1):
            _add_table_row(ldr_t, [str(idx), dr.get('device_name', '-'), dr.get('model', '-'), dr.get('ip_address', '-'), dr.get('location', '-')])
    doc.add_paragraph()
    _add_chapter_heading(doc, '八', '网络运行建议')
    doc.add_paragraph(sections.get('flood_advice', ''))
    _add_chapter_heading(doc, '九', '售后服务电话')
    doc.add_paragraph(f'运维工程师 {sections.get("tech_support", "")}')
    doc.add_paragraph(f'业务经理 {sections.get("complaint", "")}')
    doc.add_paragraph(f'业主签字：{sections.get("owner_sign", "")}')
    doc.add_paragraph('运维公司：江西丰功信息技术有限公司')
    doc.add_paragraph(f'报告生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
    seal = sections.get('seal_image', '')
    if seal and os.path.exists(seal.lstrip('/')):
        try:
            doc.add_picture(seal.lstrip('/'), width=Inches(1.5))
        except:
            pass

    # 保存
    reports_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'reports')
    os.makedirs(reports_dir, exist_ok=True)
    title_safe = inspection.title.replace('/', '_').replace('\\', '_')[:50] if inspection and inspection.title else '巡检'
    filename = f'巡检报告_{title_safe}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    filepath = os.path.join(reports_dir, filename)
    doc.save(filepath)
    return filepath


def _build_device_results_from_values(inspection, field_values):
    """从 field_values_json 和 Device 表重建 device_results"""
    from models import Device
    results = []
    if not field_values:
        return results
    for dev_key, items_dict in field_values.items():
        if not isinstance(items_dict, dict):
            continue
        device = Device.query.filter_by(device_name=dev_key).first()
        cat = _infer_device_category(device.device_type) if device else '其他'
        items_list = []
        for item_name, item_value in items_dict.items():
            if item_name == '系统运行时间':
                continue
            items_list.append({'name': item_name, 'value': str(item_value) if item_value else '', 'field_type': 'text', 'help_text': ''})
        results.append({
            'device_name': dev_key, 'device_type': device.device_type if device else '',
            'location': device.location if device else '', 'model': device.model if device else '',
            'ip_address': device.ip_address if device else '', 'os_version': device.os_version if device else '',
            'uptime': items_dict.get('系统运行时间', '-'), 'template_name': f'{cat}通用巡检',
            'template_category': cat, 'items': items_list, 'summary': '', 'photos': [],
        })
    return results


def _infer_device_category(device_type):
    dt = (device_type or '').lower()
    if any(k in dt for k in ['服务器', 'server']): return '服务器'
    if any(k in dt for k in ['交换机', '路由器', 'switch', 'router']): return '网络设备'
    if any(k in dt for k in ['防火墙', 'ips', 'ids', 'waf', 'vpn', '上网行为']): return '安全设备'
    if any(k in dt for k in ['ups', '电源']): return 'UPS'
    if any(k in dt for k in ['空调', '精密空调']): return '空调'
    return '其他'
